"""
AI Internship Finder Agent — Flask Backend
==========================================
A Flask server that powers the AI Internship Finder chatbot.
Uses LangChain + LangGraph with OpenAI to create an intelligent agent that can
search for internships based on user skills, preferences, and resume.
"""

import os
import sys
# Force UTF-8 output so emoji in print() don't crash on Windows terminals
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
import json
import tempfile
import base64
import io
import threading
from collections import OrderedDict
import requests
from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from fpdf import FPDF
import docx as docx_lib

# from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
# from langgraph.prebuilt import create_react_agent

from tools import search_internships, filter_india_jobs
from mock_interview import mock_interview_bp

# ──────────────────────────────────────────────
# Load environment variables
# ──────────────────────────────────────────────
load_dotenv()

app = Flask(
    __name__,
    template_folder="internship_ai_agent",
    static_folder="internship_ai_agent",
    static_url_path="/static"
)
CORS(app)
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024  # 32 MB max upload

import uuid
def _resolve_upload_folder():
    explicit = os.getenv("UPLOAD_DIR", "").strip()
    if explicit:
        return explicit
    if os.getenv("RENDER") and os.path.isdir("/var/data"):
        return "/var/data/uploads"
    return os.path.join(os.path.dirname(__file__), 'uploads')


UPLOAD_FOLDER = _resolve_upload_folder()
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
RESUME_VARIATION_MEMORY_PATH = os.path.join(UPLOAD_FOLDER, 'resume_variation_memory.json')
RESUME_VARIATION_LOCK = threading.RLock()
CHAT_HISTORY_LOCK = threading.RLock()
MAX_CHAT_SESSIONS = max(50, int(os.getenv("MAX_CHAT_SESSIONS", "2000")))
MAX_HISTORY_MESSAGES = max(5, int(os.getenv("MAX_HISTORY_MESSAGES", "6")))
EXTERNAL_HTTP_TIMEOUT = max(3, int(os.getenv("EXTERNAL_HTTP_TIMEOUT", "20")))

# ── SQLite Database ──
import database as db

def get_resume_stream_from_req(req):
    # 1. Direct file upload takes priority
    if 'resume' in req.files and req.files['resume'].filename != '':
        file_obj = req.files['resume']
        payload = file_obj.read()
        file_obj.seek(0)
        return io.BytesIO(payload)

    # 2. Try identifying the session_id
    session_id = req.form.get("session_id") if req.form else None
    if not session_id and req.is_json:
        session_id = req.json.get("session_id")

    # 3. Check if the session_id from request has a file
    if session_id:
        filepath = os.path.join(UPLOAD_FOLDER, f"{session_id}.pdf")
        if os.path.exists(filepath):
            with open(filepath, 'rb') as f:
                return io.BytesIO(f.read())

    # 4. If no file yet, try identifying via Auth Token (the reliable way for logged-in users)
    auth_header = req.headers.get('Authorization', '')
    if auth_header.startswith('Bearer '):
        token = auth_header.replace('Bearer ', '').strip()
        if token and token != 'null' and token != 'undefined':
            user = db.get_user_by_token(token)
            if user and user.get('session_id'):
                sid = user.get('session_id')
                filepath = os.path.join(UPLOAD_FOLDER, f"{sid}.pdf")
                if os.path.exists(filepath):
                    with open(filepath, 'rb') as f:
                        return io.BytesIO(f.read())

    return None


def _load_resume_variation_memory():
    with RESUME_VARIATION_LOCK:
        try:
            if not os.path.exists(RESUME_VARIATION_MEMORY_PATH):
                return {}
            with open(RESUME_VARIATION_MEMORY_PATH, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return data if isinstance(data, dict) else {}
        except Exception:
            return {}


def _save_resume_variation_memory(data):
    with RESUME_VARIATION_LOCK:
        try:
            temp_path = f"{RESUME_VARIATION_MEMORY_PATH}.{uuid.uuid4().hex}.tmp"
            with open(temp_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            for attempt in range(5):
                try:
                    os.replace(temp_path, RESUME_VARIATION_MEMORY_PATH)
                    break
                except PermissionError:
                    if attempt == 4:
                        raise
                    import time
                    time.sleep(0.03 * (attempt + 1))
        except Exception as ex:
            print(f"[WARN] Failed to save resume variation memory: {ex}")


def _get_resume_variation_history(user_id):
    if not user_id:
        return {'summary': [], 'experience': []}
    memory = _load_resume_variation_memory()
    entry = memory.get(str(user_id), {})
    if not isinstance(entry, dict):
        return {'summary': [], 'experience': []}
    summary = entry.get('summary', [])
    experience = entry.get('experience', [])
    if not isinstance(summary, list):
        summary = []
    if not isinstance(experience, list):
        experience = []
    return {'summary': summary, 'experience': experience}


def _remember_resume_variation(user_id, summary_text, experience_text, keep=6):
    if not user_id:
        return
    with RESUME_VARIATION_LOCK:
        try:
            if not os.path.exists(RESUME_VARIATION_MEMORY_PATH):
                memory = {}
            else:
                with open(RESUME_VARIATION_MEMORY_PATH, 'r', encoding='utf-8') as f:
                    loaded = json.load(f)
                memory = loaded if isinstance(loaded, dict) else {}
        except Exception:
            memory = {}

        entry = memory.get(str(user_id), {})
        if not isinstance(entry, dict):
            entry = {}
        summary_hist = entry.get('summary', [])
        experience_hist = entry.get('experience', [])
        if not isinstance(summary_hist, list):
            summary_hist = []
        if not isinstance(experience_hist, list):
            experience_hist = []

        if summary_text and summary_text.strip():
            summary_hist.append(summary_text.strip())
        if experience_text and experience_text.strip():
            experience_hist.append(experience_text.strip())

        entry['summary'] = summary_hist[-keep:]
        entry['experience'] = experience_hist[-keep:]
        memory[str(user_id)] = entry

        temp_path = f"{RESUME_VARIATION_MEMORY_PATH}.{uuid.uuid4().hex}.tmp"
        try:
            with open(temp_path, 'w', encoding='utf-8') as f:
                json.dump(memory, f, ensure_ascii=False, indent=2)
            for attempt in range(5):
                try:
                    os.replace(temp_path, RESUME_VARIATION_MEMORY_PATH)
                    break
                except PermissionError:
                    if attempt == 4:
                        raise
                    import time
                    time.sleep(0.03 * (attempt + 1))
        except Exception as ex:
            print(f"[WARN] Failed to save resume variation memory: {ex}")

app.register_blueprint(mock_interview_bp)

# ──────────────────────────────────────────────
# Native Auth API Routes (no Node.js needed)
# ──────────────────────────────────────────────
import hashlib
import secrets

def _hash_password(pw):
    return hashlib.sha256(pw.encode()).hexdigest()

def _make_token(user_id):
    return secrets.token_hex(32)


@app.route('/api/auth/register', methods=['POST'])
def auth_register():
    """Handle registration from the main register.html form (first_name + last_name fields)."""
    try:
        # register.html sends form-data with first_name / last_name
        first_name = request.form.get('first_name', '').strip()
        last_name  = request.form.get('last_name', '').strip()
        email      = request.form.get('email', '').strip().lower()
        password   = request.form.get('password', '')
        target_role = request.form.get('role', '').strip()
        full_name  = (first_name + ' ' + last_name).strip() or email.split('@')[0]

        if not email or not password:
            return jsonify({'status': 'error', 'error': 'Email and password are required'}), 400

        # Check if user already exists
        if db.get_user_by_email(email):
            return jsonify({'status': 'error', 'error': 'User with this email already exists'}), 409

        # Handle optional resume upload
        skills = []
        role   = target_role or 'Software Engineer'
        session_id = None
        if 'resume' in request.files:
            resume_file = request.files['resume']
            session_id  = str(uuid.uuid4())
            filepath    = os.path.join(UPLOAD_FOLDER, f"{session_id}.pdf")
            resume_file.save(filepath)
            try:
                from mock_interview import extract_text_from_pdf, extract_skills, infer_role
                with open(filepath, 'rb') as f:
                    text = extract_text_from_pdf(f)
                skills = extract_skills(text)
                if skills:
                    role = infer_role(skills)
            except Exception as ex:
                print(f"[Register] Resume parse error: {ex}")

        user = db.create_user(
            email=email,
            password_hash=_hash_password(password),
            full_name=full_name,
            target_role=role,
            session_id=session_id or '',
            extracted_skills=skills,
            extracted_role=role
        )
        if not user:
            return jsonify({'status': 'error', 'error': 'User with this email already exists'}), 409

        token = user.get('token', '')
        return jsonify({
            'status': 'success',
            'token': token,
            'session_id': session_id,
            'skills': skills,
            'role': role
        }), 201
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'status': 'error', 'error': str(e)}), 500


@app.route('/api/auth/signup', methods=['POST'])
def auth_signup():
    try:
        # Accept both JSON and form-data
        if request.is_json:
            data = request.get_json()
            email = data.get('email', '').strip().lower()
            password = data.get('password', '')
            full_name = data.get('fullName', data.get('full_name', ''))
        else:
            email = request.form.get('email', '').strip().lower()
            password = request.form.get('password', '')
            full_name = request.form.get('fullName', request.form.get('full_name', ''))

        if not email or not password:
            return jsonify({"error": "Email and password are required"}), 400

        if db.get_user_by_email(email):
            return jsonify({"error": "User with this email already exists"}), 409

        # Handle optional resume upload
        skills = []
        role = 'Software Engineer'
        session_id = None
        if 'resume' in request.files:
            resume_file = request.files['resume']
            session_id = str(uuid.uuid4())
            filepath = os.path.join(UPLOAD_FOLDER, f"{session_id}.pdf")
            resume_file.save(filepath)
            try:
                from mock_interview import extract_text_from_pdf, extract_skills, infer_role
                with open(filepath, 'rb') as f:
                    text = extract_text_from_pdf(f)
                skills = extract_skills(text)
                role = infer_role(skills) if skills else 'Software Engineer'
            except Exception as e:
                print(f"[Signup] Resume parsing error: {e}")

        user = db.create_user(
            email=email,
            password_hash=_hash_password(password),
            full_name=full_name or email.split('@')[0],
            target_role=role,
            session_id=session_id or '',
            extracted_skills=skills,
            extracted_role=role
        )
        if not user:
            return jsonify({"error": "User with this email already exists"}), 409

        safe_user = db.user_to_safe_dict(user)
        return jsonify({
            'token': user['token'],
            'session_id': session_id,
            'user': safe_user,
            'skills': skills,
            'role': role
        }), 201
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/auth/login', methods=['POST'])
def auth_login():
    try:
        if request.is_json:
            data = request.get_json()
            email = data.get('email', '').strip().lower()
            password = data.get('password', '')
        else:
            email = request.form.get('email', '').strip().lower()
            password = request.form.get('password', '')

        if not email or not password:
            return jsonify({'error': 'Email and password are required'}), 400

        user = db.get_user_by_email(email)
        if not user:
            return jsonify({'error': 'Invalid email or password'}), 401

        # Verify password
        if user.get('password_hash') != _hash_password(password):
            return jsonify({'error': 'Invalid email or password'}), 401

        # Refresh token on every login
        token = db.refresh_user_token(user['_id'])

        skills    = user.get('extractedSkills', [])
        role      = user.get('extractedRole', user.get('extracted_role', 'Software Engineer'))
        full_name = user.get('fullName', user.get('full_name', '')) or email.split('@')[0]

        safe_user = db.user_to_safe_dict(user)
        safe_user['token'] = token  # Use the refreshed token
        return jsonify({'token': token, 'user': safe_user, 'skills': skills, 'role': role})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/auth/upload-resume', methods=['POST'])
def auth_upload_resume():
    try:
        auth_header = request.headers.get('Authorization', '')
        token = auth_header.replace('Bearer ', '').strip()

        user = db.get_user_by_token(token)
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401

        if 'resume' not in request.files:
            return jsonify({'error': 'No resume file provided'}), 400

        resume_file = request.files['resume']
        session_id = str(uuid.uuid4())
        filepath = os.path.join(UPLOAD_FOLDER, f"{session_id}.pdf")
        resume_file.save(filepath)

        skills = []
        role = 'Software Engineer'
        try:
            from mock_interview import extract_text_from_pdf, extract_skills, infer_role
            with open(filepath, 'rb') as f:
                text = extract_text_from_pdf(f)
            skills = extract_skills(text)
            role = infer_role(skills) if skills else 'Software Engineer'
        except Exception as e:
            print(f"[UploadResume] Error: {e}")

        db.update_user(user['_id'],
            extracted_skills=skills,
            extracted_role=role,
            session_id=session_id
        )

        # Re-fetch updated user
        user = db.get_user_by_id(user['_id'])
        safe_user = db.user_to_safe_dict(user)
        return jsonify({'user': safe_user, 'skills': skills, 'role': role})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ── In-memory reset token store: {email: {code, expires_at}} ──
import time as _time
_reset_tokens = {}

@app.route('/forgot-password')
def forgot_password_page():
    return render_template('forgot-password.html')

@app.route('/api/auth/forgot-password', methods=['POST'])
def auth_forgot_password():
    try:
        email = request.form.get('email', '').strip().lower()
        if not email:
            return jsonify({'error': 'Email is required'}), 400

        user = db.get_user_by_email(email)
        if not user:
            return jsonify({'error': 'No account found with that email address.'}), 404

        # Generate a 6-digit OTP
        import random as _random
        code = str(_random.randint(100000, 999999))
        _reset_tokens[email] = {
            'code': code,
            'expires_at': _time.time() + 600  # 10 minutes
        }
        print(f"[Password Reset] Code for {email}: {code}")

        # Return the code directly (local/demo mode — no email server)
        return jsonify({'status': 'ok', 'code': code,
                        'message': 'Reset code generated. Copy it from the box above.'})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/auth/reset-password', methods=['POST'])
def auth_reset_password():
    try:
        email    = request.form.get('email', '').strip().lower()
        code     = request.form.get('code', '').strip()
        new_pw   = request.form.get('new_password', '')

        if not email or not code or not new_pw:
            return jsonify({'error': 'Email, code and new password are required.'}), 400
        if len(new_pw) < 6:
            return jsonify({'error': 'Password must be at least 6 characters.'}), 400

        token_data = _reset_tokens.get(email)
        if not token_data:
            return jsonify({'error': 'No reset code found. Please request a new one.'}), 400
        if _time.time() > token_data['expires_at']:
            _reset_tokens.pop(email, None)
            return jsonify({'error': 'Reset code has expired. Please request a new one.'}), 400
        if token_data['code'] != code:
            return jsonify({'error': 'Invalid reset code. Check and try again.'}), 400

        # All good — update the password
        user = db.get_user_by_email(email)
        if not user:
            return jsonify({'error': 'Account not found.'}), 404

        db.update_user_password(email, _hash_password(new_pw))
        _reset_tokens.pop(email, None)       # invalidate used token

        return jsonify({'status': 'success', 'message': 'Password reset successfully.'})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ──────────────────────────────────────────────
# Profile & Interview History API Routes
# ──────────────────────────────────────────────

def _get_user_from_request():
    """Extract the authenticated user from the request (via Authorization header or session_id)."""
    auth_header = request.headers.get('Authorization', '')
    token = auth_header.replace('Bearer ', '').strip()
    if token:
        return db.get_user_by_token(token)
    return None


@app.route('/api/profile', methods=['GET'])
def api_get_profile():
    """Fetch the authenticated user's profile from the database."""
    try:
        user = _get_user_from_request()
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401

        profile = db.get_profile(user['_id'])
        # Merge user-level fields into profile for convenience
        profile['fullName'] = user.get('fullName', user.get('full_name', ''))
        profile['email'] = user.get('email', '')
        profile['role'] = user.get('extractedRole', user.get('extracted_role', ''))
        profile['extractedSkills'] = user.get('extractedSkills', [])
        return jsonify({'profile': profile})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/profile', methods=['POST'])
def api_save_profile():
    """Save or update the authenticated user's profile."""
    try:
        user = _get_user_from_request()
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401

        if request.is_json:
            data = request.get_json()
        else:
            data = request.form.to_dict()

        # Update the user's name if provided
        name = data.get('name', data.get('fullName', ''))
        if name:
            db.update_user(user['_id'], full_name=name)

        # Save profile fields
        db.save_profile(user['_id'], data)

        return jsonify({'status': 'success', 'message': 'Profile saved'})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/interview-history', methods=['GET'])
def api_get_interview_history():
    """Fetch the authenticated user's interview history."""
    try:
        user = _get_user_from_request()
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401

        history = db.get_interview_history(user['_id'])
        return jsonify({'history': history})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/interview-history', methods=['POST'])
def api_save_interview_history():
    """Save an interview session result."""
    try:
        user = _get_user_from_request()
        user_id = user['_id'] if user else None

        data = request.get_json() or {}
        session_id = db.save_interview_session(user_id, data)

        return jsonify({'status': 'success', 'session_id': session_id})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ──────────────────────────────────────────────
# LangChain Agent Setup
# ──────────────────────────────────────────────

# ── LLM Setup: Google Gemini 2.5 Flash (free tier) primary, Ollama fallback ──
llm = None
llm_resume = None

def setup_llm():
    global llm, llm_resume
    GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "")
    if GOOGLE_API_KEY:
        from langchain_google_genai import ChatGoogleGenerativeAI
        llm = ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            google_api_key=GOOGLE_API_KEY,
            temperature=0.4,
            convert_system_message_to_human=True,
            top_p=0.95,
            top_k=40
        )
        llm_resume = ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            google_api_key=GOOGLE_API_KEY,
            temperature=0.8,
            convert_system_message_to_human=True,
            top_p=0.98,
            top_k=50
        )
        print("[OK] LLM: Google Gemini 2.5 Flash")
    else:
        from langchain_ollama import ChatOllama
        llm = ChatOllama(model="llama3.2:1b", temperature=0.1)
        llm_resume = llm
        print("[WARN] GOOGLE_API_KEY not set - falling back to local Ollama llama3.2:1b")
# Tools available to the agent
tools = [search_internships]

# System prompt for the AI agent (Optimized for speed & anti-hallucination)
SYSTEM_PROMPT = """You are an AI Internship Finder Agent. Your goal is to help students find internships efficiently.

RULES:
1. ALWAYS use the `search_internships` tool to find internships.
2. The `search_internships` tool will return a fully formatted ```internship_cards``` block. YOU MUST ECHO THIS EXACT BLOCK TO THE USER VERBATIM!
3. DO NOT change the JSON. DO NOT change the apply_links. DO NOT invent URLs. DO NOT hallucinate.
4. Keep all conversational text extremely concise. Answer in 1 short sentence.
"""

# Create the agent using LangGraph's create_react_agent
agent = None
def setup_agent():
    global agent
    if llm:
        from langgraph.prebuilt import create_react_agent
        agent = create_react_agent(llm, tools)

# ──────────────────────────────────────────────
# In-memory chat history (per session — resets on server restart)
# ──────────────────────────────────────────────
chat_histories = OrderedDict()


def get_chat_history(session_id: str) -> list:
    """Get or create chat history for a session."""
    sid = str(session_id or "default")
    with CHAT_HISTORY_LOCK:
        history = chat_histories.get(sid)
        if history is None:
            from langchain_core.messages import SystemMessage
            history = [SystemMessage(content=SYSTEM_PROMPT)]
            chat_histories[sid] = history
        else:
            chat_histories.move_to_end(sid)

        while len(chat_histories) > MAX_CHAT_SESSIONS:
            chat_histories.popitem(last=False)
        return history


def append_chat_messages(session_id: str, messages):
    """Append messages to chat history and trim to bounded length."""
    if not isinstance(messages, list):
        messages = [messages]
    with CHAT_HISTORY_LOCK:
        history = get_chat_history(session_id)
        history.extend(messages)
        if len(history) > MAX_HISTORY_MESSAGES:
            history[:] = [history[0]] + history[-(MAX_HISTORY_MESSAGES - 1):]


def get_chat_history_snapshot(session_id: str) -> list:
    """Return a shallow copy so model calls aren't affected by concurrent writes."""
    with CHAT_HISTORY_LOCK:
        return list(get_chat_history(session_id))


# ──────────────────────────────────────────────
# Routes
# ──────────────────────────────────────────────

@app.route("/")
def index():
    """Serve the premium liquid glass landing page."""
    return render_template("landing.html")

@app.route("/about-us")
def about_us():
    """Serve the about us page."""
    return render_template("about-us.html")

@app.route("/privacy-policy")
def privacy_policy():
    """Serve the privacy policy page."""
    return render_template("privacy-policy.html")

@app.route("/terms-and-conditions")
def terms_and_conditions():
    """Serve the terms and conditions page."""
    return render_template("terms-and-conditions.html")

@app.route("/contact-us")
def contact_us():
    """Serve the contact us page."""
    return render_template("contact-us.html")

@app.route("/profile")
def profile():
    """Serve the user profile page."""
    return render_template("profile.html")

@app.route("/images/<path:filename>")
def serve_image(filename):
    """Serve images from the internship_ai_agent/images folder."""
    from flask import send_from_directory
    images_folder = os.path.join(os.path.dirname(__file__), 'internship_ai_agent', 'images')
    return send_from_directory(images_folder, filename)

@app.route("/user-dashboard")
def user_dashboard():
    """Serve the post-auth user dashboard."""
    return render_template("user-dashboard.html")


@app.route("/dashboard")
def dashboard():
    """Serve the new professional dashboard."""
    return render_template("user-dashboard.html")

@app.route("/bot")
def bot():
    """Serve the original interactive chatbot UI."""
    return render_template("bot.html")

@app.route('/login')
def login():
    return render_template('login.html')

@app.route('/register')
def register():
    return render_template('register.html')

# ── Auth routes are now proxied to Node.js backend via /api/auth/<path:subpath> ──

@app.route('/api/generate-resume', methods=['POST'])
def generate_resume():
    try:
        user = _get_user_from_request()
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401

        data = request.get_json() or {}
        variation_nonce = uuid.uuid4().hex[:10]
        variation_history = _get_resume_variation_history(user.get('_id'))

        def build_avoid_phrases(history_items, limit=8):
            import re
            phrases = []
            for item in history_items[-3:]:
                if not item or not isinstance(item, str):
                    continue
                for chunk in re.split(r'[.\n;]+', item):
                    cleaned = " ".join(chunk.strip().split())
                    if len(cleaned.split()) >= 4:
                        phrases.append(cleaned)
            unique = []
            seen = set()
            for phrase in phrases:
                key = phrase.lower()
                if key in seen:
                    continue
                seen.add(key)
                unique.append(phrase[:120])
                if len(unique) >= limit:
                    break
            return unique

        summary_avoid_phrases = build_avoid_phrases(variation_history.get('summary', []))
        experience_avoid_phrases = build_avoid_phrases(variation_history.get('experience', []))

        def format_avoid_phrases(phrases):
            if not phrases:
                return "None"
            return "\n".join(f'- "{p}"' for p in phrases)

        def pick_non_repeating(options, context_key):
            import random
            if not options:
                return ""
            history = variation_history.get(context_key, [])
            previous = history[-1].strip().lower() if history else ""
            filtered = [opt for opt in options if opt and opt.strip().lower() != previous]

            avoid_phrases = summary_avoid_phrases if context_key == 'summary' else experience_avoid_phrases
            phrase_filtered = [
                opt for opt in filtered
                if not any(p.lower() in opt.lower() for p in avoid_phrases[:3])
            ]
            pool = phrase_filtered or filtered or options
            return random.choice(pool)
        
        # Professional FPDF Generation
        class ProResumePDF(FPDF):
            def header(self):
                pass
            def footer(self):
                self.set_y(-15)
                self.set_font('Arial', 'I', 8)
                self.set_text_color(150, 150, 150)
                self.cell(0, 10, 'Generated by InternAI Pro', 0, 0, 'C')

        pdf = ProResumePDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Colors
        accent_r, accent_g, accent_b = 33, 50, 94 # Dark Slate Blue
        text_r, text_g, text_b = 50, 50, 50
        
        # 1. HEADER SECTON (Colored background)
        pdf.set_fill_color(accent_r, accent_g, accent_b)
        pdf.rect(0, 0, 210, 45, 'F')
        
        # Name
        pdf.set_y(10)
        pdf.set_font('Arial', 'B', 26)
        pdf.set_text_color(255, 255, 255)
        name = f"{data.get('firstName', '')} {data.get('lastName', '')}".strip()
        if not name: name = user.get('full_name') or 'Professional Name'
        pdf.cell(0, 10, txt=name.upper(), ln=True, align='C')
        
        # Role
        pdf.set_font('Arial', '', 14)
        pdf.set_text_color(200, 220, 255)
        pdf.cell(0, 8, txt=data.get('profession', 'Software Engineer').upper(), ln=True, align='C')
        
        # Contact Info
        pdf.set_font('Arial', '', 10)
        pdf.set_text_color(255, 255, 255)
        contact = f"{data.get('location', '')}  |  {data.get('phone', '')}  |  {data.get('email', '')}"
        pdf.cell(0, 6, txt=contact, ln=True, align='C')
        
        pdf.set_y(55) # Move below header
        
        # Helper function for sections
        def add_section_title(title):
            pdf.ln(4)
            pdf.set_font('Arial', 'B', 14)
            pdf.set_text_color(accent_r, accent_g, accent_b)
            pdf.cell(0, 8, txt=title.upper(), ln=True)
            # Horizontal line
            y = pdf.get_y()
            pdf.set_draw_color(accent_r, accent_g, accent_b)
            pdf.line(10, y, 200, y)
            pdf.ln(3)

        # AI Enhancement of Text using Gemini 2.5 Flash with specialized prompts
        def enhance_text(text, context, user_info=None):
            """Aggressive AI enhancement with VARIED output each time."""
            from langchain_core.messages import HumanMessage
            import random
            
            user_info = user_info or {}
            profession = user_info.get('profession', 'Software Engineer')
            skills = user_info.get('skills', '')
            years_exp = user_info.get('years_exp', '2')
            
            # If text is empty or very minimal, generate from scratch
            if not text or len(text.strip()) < 5:
                if context == "professional summary":
                    return generate_professional_summary(user_info)
                elif context == "work experience description/achievements":
                    return generate_job_description(user_info)
                return text
            
            # VARIED PROMPTS - Different focus each time for unique output
            if context == "professional summary":
                variation_styles = [
                    {
                        "name": "Impact-Driven",
                        "start_words": ["Results-oriented", "Impact-driven", "Proven"],
                        "focus": "business outcomes and measurable results"
                    },
                    {
                        "name": "Innovation-Focused",
                        "start_words": ["Innovative", "Forward-thinking", "Strategic"],
                        "focus": "technical innovation and cutting-edge solutions"
                    },
                    {
                        "name": "Leadership-Oriented",
                        "start_words": ["Dynamic", "Visionary", "Driven"],
                        "focus": "team leadership and collaborative achievements"
                    }
                ]
                
                style = random.choice(variation_styles)
                start_word = random.choice(style["start_words"])
                
                prompt = f"""REWRITE this into a PROFESSIONAL resume summary with focus on {style["focus"]}.

INPUT: {text}
VARIATION_TOKEN: {variation_nonce}
RECENT_PHRASES_TO_AVOID:
{format_avoid_phrases(summary_avoid_phrases)}

STYLE: {style["name"]}
- Start with: "{start_word}"
- Emphasize: {style["focus"]}

REQUIREMENTS:
1. Exactly 3-4 sentences (different order/structure than previous versions)
2. MUST start with: "{start_word}"
3. Include specific years or domain expertise
4. Use strong verbs: Engineered, Led, Delivered, Architected, Optimized, Transformed
5. Show concrete business impact
6. NO personal pronouns (I, me, we)
7. NO generic phrases
8. Avoid reusing wording from RECENT_PHRASES_TO_AVOID

VARY THE STRUCTURE - Make it unique:
- Option A: [Adjective] [Profession] → Specialized in → Proven track record of → Dedicated to
- Option B: [Adjective] [Professional] → Track record of → Deep expertise in → Passionate about
- Option C: [Adjective] [Role] → Demonstrated success in → Technical excellence in → Committed to

Pick a DIFFERENT structure than usual. Output ONLY the summary - NO PREAMBLE."""

            elif context == "work experience description/achievements":
                variation_styles = [
                    {
                        "name": "Technical Excellence",
                        "focus": "technical depth and architectural decisions",
                        "lead_verbs": ["Engineered", "Architected", "Designed", "Implemented"]
                    },
                    {
                        "name": "Business Impact",
                        "focus": "revenue, cost savings, and business metrics",
                        "lead_verbs": ["Delivered", "Optimized", "Accelerated", "Transformed"]
                    },
                    {
                        "name": "Leadership & Scale",
                        "focus": "team leadership and scaling systems",
                        "lead_verbs": ["Led", "Spearheaded", "Pioneered", "Established"]
                    }
                ]
                
                style = random.choice(variation_styles)
                selected_verbs = ", ".join(random.sample(style["lead_verbs"], 2))
                
                prompt = f"""Transform this into PROFESSIONAL achievement bullets with focus on {style["focus"]}.

INPUT: {text}
VARIATION_TOKEN: {variation_nonce}
RECENT_PHRASES_TO_AVOID:
{format_avoid_phrases(experience_avoid_phrases)}
STYLE: {style["name"]}
PRIMARY VERBS: {selected_verbs}

REQUIREMENTS FOR EACH BULLET:
1. Start with: {selected_verbs} (or similar power verbs)
2. MUST have: Specific metrics/numbers with context (%, time, $, users, improvement)
3. Include: Technology/tool names
4. Format: [Verb] [What] using [Tech], [Result with number]
5. One line, scannable
6. NO gerunds, NO personal pronouns
7. Avoid reusing wording from RECENT_PHRASES_TO_AVOID

CREATE 4 VARIED BULLETS showing:
- Different metrics (% improvement, time saved, volume processed, cost reduction)
- Different technologies mentioned
- Different business/technical outcomes
- Different team/scope sizes

Examples of VARIED formats:
• "Engineered X using Y, improving Z by 45%" 
• "Led team of X, architecting Y that handles Z daily"
• "Optimized X, reducing Z from Y to Y (60% improvement)"
• "Delivered X using Y, enabling Z outcome"

Output ONLY bullets on separate lines starting with "- ", NO PREAMBLE."""


            else:
                prompt = f"Rewrite this professionally for a resume: {text}"
            
            try:
                print(f"[DEBUG] Calling LLM for {context} enhancement...")
                res = llm_resume.invoke([HumanMessage(content=prompt)])
                result = res.content.strip()
                
                # Clean up formatting
                result = result.replace('```', '').replace('**', '').strip()
                
                # Validate output
                if result and len(result.strip()) > 10:
                    print(f"[SUCCESS] Enhanced {context}: {result[:80]}...")
                    return result
                else:
                    print(f"[WARN] LLM returned empty result for {context}, using fallback")
                    return fallback_enhance(text, context, user_info)
                    
            except Exception as e:
                print(f"[ERROR] LLM Enhancement failed ({context}): {e}")
                return fallback_enhance(text, context, user_info)
        
        def fallback_enhance(text, context, user_info):
            """Fallback enhancement with VARIED templates - never the same twice."""
            
            user_info = user_info or {}
            profession = user_info.get('profession', 'Software Engineer')
            skills = user_info.get('skills', '')
            years_exp = user_info.get('years_exp', '2')
            
            if context == "professional summary":
                kw_list = [s.strip() for s in skills.split(',')[:3]] if skills else ['modern technologies', 'scalable systems', 'innovative solutions']
                has_years = any(str(i) in text for i in range(1, 50))
                years = years_exp if has_years else "2+"
                
                # Multiple varied templates so fallback is never identical
                templates = [
                    f"Results-oriented {profession} with {years} years of proven track record in delivering impactful solutions. Deep expertise in {kw_list[0]} combined with strong technical foundation. Committed to driving innovation and leveraging technology to solve complex business challenges.",
                    
                    f"Driven {profession} with {years}+ years of hands-on experience architecting and delivering scalable systems. Specialized in {kw_list[0]} with demonstrated ability to optimize performance and drive measurable outcomes. Passionate about continuous learning and collaborating across teams to achieve organizational goals.",
                    
                    f"Innovative {profession} with {years} years of proven expertise in building solutions using {kw_list[0]}. Strong track record of translating complex requirements into technical excellence. Dedicated to leveraging best practices and emerging technologies to deliver competitive advantage and business value.",
                    
                    f"Strategic {profession} with {years} years of experience delivering high-impact solutions. Specialized in {kw_list[0]} with proven ability to optimize systems and enhance operational efficiency. Committed to fostering collaborative environments and driving technical excellence across projects."
                ]
                return pick_non_repeating(templates, 'summary')
            
            elif context == "work experience description/achievements":
                kw_list = [s.strip() for s in skills.split(',')[:3]] if skills else ['technical solutions', 'modern frameworks', 'scalable systems']
                
                # Multiple varied bullet templates
                template_sets = [
                    [
                        f"- Engineered robust {kw_list[0]} using modern technologies, improving system performance by 35%+",
                        "- Led cross-functional initiatives resulting in 40% efficiency gain across team productivity",
                        f"- Architected and implemented {kw_list[1]} reducing infrastructure costs by 25%",
                        "- Mentored junior developers and established code quality standards improving delivery velocity"
                    ],
                    [
                        f"- Delivered enterprise-grade {kw_list[0]} handling 10K+ concurrent users with 99.9% availability",
                        "- Optimized database operations and API performance, reducing latency by 60%",
                        f"- Spearheaded migration to {kw_list[1]}, enabling faster feature releases and deployments",
                        "- Collaborated with product teams to define technical roadmap and deliver on aggressive timelines"
                    ],
                    [
                        f"- Architected {kw_list[0]} using {kw_list[1]}, increasing throughput by 50% and reducing costs 30%",
                        "- Led team of 3+ engineers developing microservices handling millions of requests daily",
                        "- Implemented comprehensive testing and CI/CD pipelines reducing deployment time by 75%",
                        "- Drove technical innovation initiatives resulting in improved code maintainability and team efficiency"
                    ],
                    [
                        f"- Engineered high-performance {kw_list[0]} improving user experience and engagement by 45%",
                        "- Optimized cloud infrastructure reducing monthly operational expenses by $40K+",
                        f"- Led redesign of {kw_list[1]} architecture improving scalability for 5x user growth",
                        "- Established best practices and mentoring program improving team capabilities and retention"
                    ]
                ]
                chosen_set = pick_non_repeating(["\n".join(s) for s in template_sets], 'experience')
                return chosen_set
            
            return text
        
        def generate_professional_summary(user_info):
            """Generate compelling professional summary - VARIED each time."""
            from langchain_core.messages import HumanMessage
            import random
            
            profession = user_info.get('profession', 'Software Engineer')
            skills = user_info.get('skills', '')
            location = user_info.get('location', '')
            years_exp = user_info.get('years_exp', '2')
            
            # Top skills for context
            top_skills = [s.strip() for s in skills.split(',')[:4]] if skills else ['technology', 'problem-solving', 'collaboration']
            
            # Vary the prompt each time
            variation = random.choice([
                {
                    "angle": "business impact angle",
                    "emphasis": "measurable business outcomes and ROI"
                },
                {
                    "angle": "technical expertise angle", 
                    "emphasis": "architectural excellence and technical depth"
                },
                {
                    "angle": "innovation angle",
                    "emphasis": "forward-thinking solutions and emerging technologies"
                }
            ])
            
            prompt = f"""Generate an EXCEPTIONAL professional summary using the {variation["angle"]}.

CONTEXT:
- Position: {profession}
- Years: {years_exp} years
- Top Skills: {', '.join(top_skills)}
- Emphasis: {variation["emphasis"]}
- Variation Token: {variation_nonce}
- Recent Phrases To Avoid:
{format_avoid_phrases(summary_avoid_phrases)}

MUST CREATE 3-4 SENTENCES with unique phrasing:
1. Start with a strong adjective (different each time - Results-driven, Innovative, Proven, Strategic, Dynamic, etc.)
2. Include years and specific domain expertise
3. Highlight 2-3 key competencies showing {variation["emphasis"]}
4. Use strong action verbs (Spearheaded, Architected, Delivered, Engineered, Optimized)
5. Show concrete impact and business value
6. Unique ending that varies structure
7. Avoid wording used in Recent Phrases To Avoid

NO: Personal pronouns, generic phrases, weak verbs
OUTPUT ONLY THE SUMMARY - NO PREAMBLE."""
            
            try:
                print(f"[DEBUG] Generating professional summary ({variation['angle']})...")
                res = llm_resume.invoke([HumanMessage(content=prompt)])
                result = res.content.strip().replace('```', '').replace('**', '')
                
                if result and len(result.strip()) > 30:
                    print(f"[SUCCESS] Generated summary: {result[:80]}...")
                    return result
                    
            except Exception as e:
                print(f"[WARN] Summary generation failed: {e}")
            
            # Vary fallback templates
            primary_skill = top_skills[0] if top_skills else "technology"
            fallback_templates = [
                f"Results-driven {profession} with {years_exp}+ years of proven expertise in delivering scalable solutions. Specialized in {primary_skill} with demonstrated ability to optimize systems and drive measurable business outcomes. Committed to continuous learning and leveraging innovative technologies to solve complex challenges.",
                
                f"Strategic {profession} with {years_exp}+ years of hands-on experience architecting enterprise-scale solutions. Deep expertise in {primary_skill} and proven track record of translating business requirements into technical excellence. Passionate about innovation and delivering sustainable competitive advantage.",
                
                f"Innovative {profession} with {years_exp}+ years driving technical transformation and organizational growth. Strong foundation in {primary_skill} combined with strategic mindset for solving complex problems. Committed to fostering collaboration and mentoring teams to achieve exceptional results."
            ]
            return pick_non_repeating(fallback_templates, 'summary')
        
        def generate_job_description(user_info):
            """Generate high-impact achievement bullets - VARIED each time."""
            from langchain_core.messages import HumanMessage
            import random
            
            job_title = user_info.get('jobTitle', 'Software Engineer')
            employer = user_info.get('employer', 'Company')
            skills = user_info.get('skills', '')
            profession = user_info.get('profession', 'Software Engineer')
            
            # Extract top technologies
            tech_list = [s.strip() for s in skills.split(',')[:3]] if skills else []
            tech_str = ' and '.join(tech_list) if tech_list else 'modern technologies'
            
            # Vary the focus each time
            focus_angle = random.choice([
                "technical depth and architectural excellence",
                "business metrics and ROI (cost savings, revenue)",
                "team leadership and organizational impact",
                "scale and performance optimization"
            ])
            
            prompt = f"""Generate 4 EXCEPTIONAL achievement bullets with focus on {focus_angle}.

CONTEXT:
- Job Title: {job_title}
- Company: {employer}
- Tech Stack: {tech_str}
- Focus: {focus_angle}
- Variation Token: {variation_nonce}
- Recent Phrases To Avoid:
{format_avoid_phrases(experience_avoid_phrases)}

EACH BULLET MUST:
1. Start with DIFFERENT power verb: Engineered, Architected, Led, Optimized, Delivered, Spearheaded, Implemented, Designed
2. Include SPECIFIC metrics: percentages, cost savings, time reductions, scale metrics, or quantifiable improvements
3. Include technology name
4. Different metric types across bullets (percentage, money, time, and scale)
5. One line, professional, no personal pronouns
6. Avoid wording used in Recent Phrases To Avoid

STRUCTURE VARIATIONS:
- Verb + detail using tech, metric result
- Verb + achievement, metric plus context
- Verb + team or scope, resulting in metric
- Verb + initiative, reducing or improving metric

OUTPUT - 4 varied bullets starting with dash, NO PREAMBLE."""
            
            try:
                print(f"[DEBUG] Generating job description ({focus_angle})...")
                res = llm_resume.invoke([HumanMessage(content=prompt)])
                result = res.content.strip()
                
                # Ensure we have bullet format
                if '-' not in result:
                    result = '\n'.join(f"- {line.strip()}" for line in result.split('\n') if line.strip())
                
                if result and len(result.strip()) > 20:
                    print(f"[SUCCESS] Generated bullets: {result[:80]}...")
                    return result
                    
            except Exception as e:
                print(f"[WARN] Job description generation failed: {e}")
            
            # Vary fallback bullets
            tech = tech_list[0] if tech_list else 'modern technologies'
            fallback_options = [
                f"- Engineered scalable {tech} solutions improving system uptime from 95% to 99.9%\n" +
                "- Led migration project enabling 3x user capacity with 40% infrastructure cost reduction\n" +
                "- Optimized API response times by 55%, enhancing user experience and retention rates\n" +
                "- Implemented automated testing reducing production bugs by 65% and deployment time by 75%",
                
                f"- Architected microservices using {tech} processing 5M+ daily transactions with under 100ms latency\n" +
                "- Spearheaded performance optimization initiative reducing cloud costs by 60K annually\n" +
                "- Led team of 4 engineers delivering 8 features on schedule improving customer satisfaction by 35%\n" +
                "- Designed CI/CD pipeline enabling 50+ deployments weekly with zero downtime releases",
                
                f"- Delivered {tech} platform handling 50K concurrent users improving application stability by 45%\n" +
                "- Optimized database queries and caching strategy reducing average load time from 4s to 0.8s\n" +
                "- Led architecture redesign supporting 10x business growth and expanding market opportunities\n" +
                "- Implemented comprehensive monitoring and alerting reducing incident resolution time by 70%",
                
                f"- Engineered high-throughput {tech} system processing 100K+ records daily with 99.99% reliability\n" +
                "- Led cross-team initiative reducing operational overhead by 35% and improving code maintainability\n" +
                "- Architected event-driven architecture enabling real-time analytics reducing latency by 80%\n" +
                "- Mentored 2 junior developers improving code quality metrics by 40% and team productivity"
            ]
            return pick_non_repeating(fallback_options, 'experience')

        # Prepare user context for AI enhancement
        user_context = {
            'profession': data.get('profession', 'Software Engineer'),
            'skills': data.get('skills', ''),
            'jobTitle': data.get('jobTitle', ''),
            'employer': data.get('employer', ''),
            'location': data.get('location', '')
        }

        summary_text = enhance_text(data.get('summary', ''), "professional summary", user_context)
        job_desc_text = enhance_text(data.get('jobDesc', ''), "work experience description/achievements", user_context)

        # 2. PROFESSIONAL SUMMARY
        if summary_text:
            add_section_title("Professional Summary")
            pdf.set_font('Arial', '', 11)
            pdf.multi_cell(0, 6, txt=summary_text)
            
        # 3. EXPERIENCE
        if data.get('jobTitle') or data.get('employer'):
            add_section_title("Work Experience")
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(100, 6, txt=data.get('jobTitle', 'Role'), ln=0)
            
            # Dates aligned right
            pdf.set_font('Arial', 'I', 11)
            dates = f"{data.get('jobStart', '')} - {data.get('jobEnd', '')}"
            pdf.cell(0, 6, txt=dates, ln=True, align='R')
            
            # Employer
            pdf.set_font('Arial', 'I', 11)
            pdf.cell(0, 6, txt=data.get('employer', 'Company'), ln=True)
            
            # Description
            pdf.set_font('Arial', '', 11)
            pdf.multi_cell(0, 6, txt=job_desc_text)

        # 4. EDUCATION
        if data.get('school') or data.get('degree'):
            add_section_title("Education")
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(100, 6, txt=data.get('school', 'University'), ln=0)
            
            pdf.set_font('Arial', 'I', 11)
            pdf.cell(0, 6, txt=data.get('gradYear', ''), ln=True, align='R')
            
            pdf.set_font('Arial', '', 11)
            deg_text = f"{data.get('degree', '')} - {data.get('schoolLoc', '')}"
            pdf.cell(0, 6, txt=deg_text, ln=True)

        # 5. SKILLS
        if data.get('skills'):
            add_section_title("Skills & Competencies")
            pdf.set_font('Arial', '', 11)
            pdf.multi_cell(0, 6, txt=data.get('skills'))
            
        session_id = str(uuid.uuid4())
        filepath = os.path.join(UPLOAD_FOLDER, f"{session_id}.pdf")
        pdf.output(filepath)

        _remember_resume_variation(user.get('_id'), summary_text, job_desc_text)

        # Extract skills for DB
        from mock_interview import extract_skills
        combined_text = f"{data.get('profession', '')} {data.get('summary', '')} {data.get('jobDesc', '')} {data.get('degree', '')} {data.get('skills', '')}"
        extracted = extract_skills(combined_text)
        
        # Update user session
        db.update_user(user['_id'], session_id=session_id, extracted_skills=extracted)
        
        return jsonify({'status': 'success', 'session_id': session_id, 'skills': extracted})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/download-resume', methods=['GET'])
def download_resume():
    try:
        token = request.args.get('token')
        if not token:
            return "Unauthorized", 401
        user = db.get_user_by_token(token)
        if not user or not user.get('session_id'):
            return "No resume found", 404
        
        session_id = user.get('session_id')
        filepath = os.path.join(UPLOAD_FOLDER, f"{session_id}.pdf")
        if os.path.exists(filepath):
            from flask import send_file
            return send_file(filepath, as_attachment=True, download_name="My_AI_Resume.pdf", mimetype='application/pdf')
        return "File not found", 404
    except Exception as e:
        return str(e), 500

@app.route('/api/dashboard/analyze-resume', methods=['POST'])
def analyze_resume_for_dashboard():
    """
    Full pipeline:
      1. Extract text from PDF using pdfminer.six / PyPDF2 / pytesseract OCR
      2. Extract skills from text using the curated SKILLS_DB
      3. Scrape real internship listings from the internet for each top skill
      4. Score each listing against user's skill profile
      5. Return stats + top matched internships
    """
    from mock_interview import extract_text_from_pdf, extract_skills, infer_role, extract_resume_profile
    import requests
    from bs4 import BeautifulSoup

    stream = get_resume_stream_from_req(request)
    if not stream:
        return jsonify({"error": "No resume found. Please provide session_id or file."}), 400

    try:
        # ── Step 1: Extract text ──────────────────────────────────────
        try:
            text = extract_text_from_pdf(stream)
        except Exception as e:
            return jsonify({"error": f"Text extraction failed: {str(e)}"}), 500

        if not text or len(text.strip()) < 30:
            # Fallback: Try to use existing profile data or infer from user metadata if extraction fails
            auth_header = request.headers.get('Authorization', '')
            if auth_header.startswith('Bearer '):
                token = auth_header.replace('Bearer ', '').strip()
                user = db.get_user_by_token(token)
                if user:
                    saved_skills = user.get('extractedSkills', [])
                    if not saved_skills:
                        # If no skills saved, infer from target_role (case-insensitive)
                        from mock_interview import ROLE_MAP
                        target = user.get('target_role', 'Software Engineer')
                        # Try exact match, then try title-case
                        saved_skills = ROLE_MAP.get(target) or ROLE_MAP.get(target.title()) or ['Python', 'Communication', 'Problem Solving']
                    
                    text = f"Profile for {user.get('full_name') or 'User'}. Target Role: {user.get('target_role') or 'Software Engineer'}. Skills: {', '.join(saved_skills)}"
                else:
                    return jsonify({"error": "Could not extract readable text from this PDF and no saved profile found. Please try a text-based PDF."}), 400
            else:
                return jsonify({"error": "Could not extract readable text from this PDF. Please try a text-based PDF."}), 400

        # ── Step 2: Extract skills, role, and profile details ────────
        profile = extract_resume_profile(text)
        skills = profile.get("skills", [])
        if not skills:
            skills = ['Python', 'Communication', 'Problem Solving']

        inferred_role = profile.get("role") or infer_role(skills)

        # Update user profile and save file if provided
        auth_header = request.headers.get('Authorization', '')
        if auth_header.startswith('Bearer '):
            token = auth_header.replace('Bearer ', '').strip()
            user = db.get_user_by_token(token)
            if user:
                update_data = {
                    'extracted_skills': skills,
                    'extracted_role': inferred_role
                }
                
                # If a new resume was uploaded, save it and update session_id
                if 'resume' in request.files and request.files['resume'].filename != '':
                    import uuid
                    new_session_id = str(uuid.uuid4())
                    filepath = os.path.join(UPLOAD_FOLDER, f"{new_session_id}.pdf")
                    # Seek to beginning to ensure we save the whole file (in case extraction already read it)
                    request.files['resume'].seek(0)
                    request.files['resume'].save(filepath)
                    update_data['session_id'] = new_session_id
                
                db.update_user(user['_id'], **update_data)

        # ── Step 3: Scrape real internships ──────────────────────────
        def scrape_internships(skill_query, role_query, max_results=15):
            """Scrape real internship listings using DuckDuckGo HTML search."""
            jobs = []
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0 Safari/537.36'}

            # Force inclusion of role and intern/internship keywords
            search_terms = [
                f'"{role_query}" intern internship india 2025 2026',
                f'{skill_query} "{role_query}" internship India',
                f'site:linkedin.com "{role_query}" internship India'
            ]

            for query in search_terms:
                try:
                    search_url = f"https://html.duckduckgo.com/html/?q={requests.utils.quote(query)}"
                    resp = requests.get(search_url, headers=headers, timeout=10)
                    soup = BeautifulSoup(resp.text, 'html.parser')

                    for result in soup.select('.result')[:max_results]:
                        title_el = result.select_one('.result__title')
                        snippet_el = result.select_one('.result__snippet')
                        url_el = result.select_one('.result__url')

                        if title_el and snippet_el:
                            title = title_el.get_text(strip=True)
                            snippet = snippet_el.get_text(strip=True)
                            url = url_el.get_text(strip=True) if url_el else ''

                            # Strict validation: Must be an internship and related to the role or tech
                            combined = (title + ' ' + snippet).lower()
                            is_intern = any(w in combined for w in ['intern', 'internship', 'entry level', 'fresher', 'graduate'])
                            
                            # Filter out non-tech/irrelevant roles like "driver", "security", etc.
                            # We check if at least one word from the role query or common tech keywords exists in title
                            role_words = role_query.lower().split()
                            tech_keywords = ['software', 'developer', 'engineer', 'backend', 'frontend', 'data', 'ai', 'cloud', 'tech', 'coding', 'programming']
                            is_relevant = any(w in title.lower() for w in role_words) or any(w in title.lower() for w in tech_keywords)
                            india_markers = ['india', 'bangalore', 'bengaluru', 'hyderabad', 'pune', 'mumbai', 'delhi', 'noida', 'gurgaon', 'gurugram', '.in']
                            is_india = any(m in combined for m in india_markers) or any(m in (url or '').lower() for m in india_markers)

                            if is_intern and is_relevant and is_india:
                                jobs.append({'title': title, 'snippet': snippet, 'url': url})

                    if len(jobs) >= 5: break
                except Exception as e:
                    print(f"[Scraper] Error: {e}")

            return jobs

        # ── Step 4: Score each scraped job against skills ─────────────
        def compute_match_score(job_text, user_skills_lower):
            """Returns a match percentage (0-100) based on skill overlap."""
            job_lower = job_text.lower()
            if not user_skills_lower: return 50 # Neutral if no skills
            matched = sum(1 for skill in user_skills_lower if skill.lower() in job_lower)
            base_score = min(100, 40 + int((matched / max(len(user_skills_lower), 1)) * 60))
            return base_score

        # Prepare queries
        skills_subset = skills[:3]
        skill_query = " ".join(skills_subset)
        raw_jobs = scrape_internships(skill_query, inferred_role, max_results=15)

        user_skills_lower = [s.lower() for s in skills]

        # Score and sort
        scored_jobs = []
        seen_titles = set()
        for job in raw_jobs:
            combined_text = job['title'] + ' ' + job['snippet']
            score = compute_match_score(combined_text, user_skills_lower)

            # Give bonus points for matching role keywords in title
            if inferred_role.lower().split()[0] in job['title'].lower():
                score = min(100, score + 8)

            clean_title = job['title'][:60]
            if clean_title not in seen_titles:
                seen_titles.add(clean_title)
                scored_jobs.append({
                    'title': clean_title,
                    'snippet': job['snippet'][:120],
                    'url': 'https://' + job['url'] if job['url'] and not job['url'].startswith('http') else job['url'],
                    'matchScore': score,
                    'company': _extract_company(job['url'], job['snippet']),
                    'mode': _guess_mode(job['snippet']),
                })

        scored_jobs.sort(key=lambda j: j['matchScore'], reverse=True)
        top_jobs = scored_jobs[:10]

        # Step 5: Compute dashboard stats
        matched_count = len(scored_jobs)
        avg_score = round(sum(j['matchScore'] for j in top_jobs) / max(len(top_jobs), 1)) if top_jobs else 75
        interviews_completed = max(1, min(20, len(skills) // 3))

        # Step 6: Skill Gap Analysis
        DEMAND_SKILLS_DB = [
            'Python','JavaScript','TypeScript','React','Node.js','Java','C++','C#','Go','Rust',
            'SQL','PostgreSQL','MongoDB','Redis','MySQL','Docker','Kubernetes','AWS','GCP','Azure',
            'Terraform','CI/CD','Linux','Git','REST API','GraphQL','Microservices','FastAPI',
            'Django','Flask','Spring Boot','Machine Learning','Deep Learning','TensorFlow','PyTorch',
            'NLP','Computer Vision','Pandas','NumPy','Scikit-learn','Data Science','Tableau',
            'Power BI','React Native','Flutter','Kotlin','Swift','Android','iOS','Figma','UI/UX',
            'Agile','Scrum','Jira','DevOps','MLOps','LangChain','OpenAI','Blockchain',
            'Cybersecurity','Spark','Hadoop','Vue.js','Angular','Next.js','Excel',
        ]
        all_job_text = ' '.join((j['title'] + ' ' + j['snippet']) for j in raw_jobs)
        all_job_lower = all_job_text.lower()
        skill_freq = {}
        for sk in DEMAND_SKILLS_DB:
            cnt = all_job_lower.count(sk.lower())
            if cnt > 0:
                skill_freq[sk] = cnt
        gap_skills = []
        for sk, cnt in sorted(skill_freq.items(), key=lambda x: x[1], reverse=True):
            if sk.lower() not in user_skills_lower:
                priority = 'high' if cnt >= 3 else 'medium' if cnt >= 2 else 'low'
                gap_skills.append({'skill': sk, 'demand': cnt, 'priority': priority})
            if len(gap_skills) >= 10:
                break

        # Step 7: Live Market Trends
        ROLE_MAP = {
            'Software Engineer': ['software engineer','backend','frontend','full stack','fullstack'],
            'Data Science':      ['data science','data scientist','data analyst','analytics'],
            'Machine Learning':  ['machine learning','ml engineer','ai engineer','deep learning'],
            'Web Development':   ['web developer','react developer','frontend developer'],
            'Mobile Dev':        ['mobile','android','ios','flutter','react native'],
            'DevOps / Cloud':    ['devops','cloud engineer','kubernetes','aws','azure'],
            'UI/UX Design':      ['ui/ux','ui design','ux design','product design','figma'],
            'Cybersecurity':     ['cybersecurity','security analyst','penetration','ethical hacking'],
        }
        role_counts = {r: sum(all_job_lower.count(kw) for kw in kws) for r, kws in ROLE_MAP.items()}
        total_hits = max(sum(role_counts.values()), 1)
        market_trends = sorted(
            [{'role': r, 'count': c, 'percent': round((c / total_hits) * 100)}
             for r, c in role_counts.items() if c > 0],
            key=lambda x: x['count'], reverse=True
        )[:8]

        # Step 8: Company Culture Insights
        def _company_insight(snippet, company):
            s = (snippet + ' ' + company).lower()
            if any(w in s for w in ['google','microsoft','amazon','meta','apple','netflix','openai']):
                return {'tag': 'FAANG-tier',  'color': '#f59e0b', 'desc': 'Top-tier tech - rigorous DSA rounds, competitive pay'}
            if any(w in s for w in ['startup','seed','series a','early stage','venture']):
                return {'tag': 'Startup',     'color': '#a855f7', 'desc': 'Startup culture - high ownership, fast-paced, direct impact'}
            if any(w in s for w in ['remote','work from home','wfh','distributed']):
                return {'tag': 'Remote-First','color': '#00f0ff', 'desc': 'Remote-first team - async culture, flexible hours'}
            if any(w in s for w in ['agile','scrum','sprint','kanban']):
                return {'tag': 'Agile',       'color': '#00e4b8', 'desc': 'Agile shop - sprints, standups, iterative delivery'}
            if any(w in s for w in ['research','phd','academic','publication']):
                return {'tag': 'Research',    'color': '#4da0ff', 'desc': 'Research-oriented - cutting-edge work, papers and experiments'}
            if any(w in s for w in ['fintech','finance','bank','trading','insurance']):
                return {'tag': 'FinTech',     'color': '#f97316', 'desc': 'Finance sector - compliance-aware, high-performance systems'}
            return     {'tag': 'Tech Co.',   'color': '#6366f1', 'desc': 'Tech company - collaborative culture, learning opportunities'}

        for job in top_jobs:
            job['insight'] = _company_insight(job.get('snippet', ''), job.get('company', ''))

        # --- Fallback: static market trends when scraping returns nothing ---
        if not market_trends:
            market_trends = [
                {'role': 'Software Engineer', 'count': 45, 'percent': 35},
                {'role': 'Data Science',      'count': 26, 'percent': 20},
                {'role': 'Machine Learning',  'count': 20, 'percent': 16},
                {'role': 'Web Development',   'count': 18, 'percent': 14},
                {'role': 'Mobile Dev',        'count': 10, 'percent': 8},
                {'role': 'DevOps / Cloud',    'count': 9,  'percent': 7},
            ]

        # --- Fallback: role-based gap skills when scraping returns nothing ---
        if not gap_skills:
            ROLE_GAP_MAP = {
                'Software Engineer':   ['Docker','Kubernetes','AWS','System Design','GraphQL','CI/CD'],
                'Frontend Developer':  ['TypeScript','Next.js','GraphQL','Jest','Docker','AWS'],
                'Data Scientist':      ['PyTorch','TensorFlow','Spark','Airflow','MLOps','Docker'],
                'Machine Learning':    ['PyTorch','MLOps','Docker','Kubernetes','Spark','GCP'],
                'Full Stack Developer':['Docker','AWS','GraphQL','Redis','Kubernetes','CI/CD'],
                'Mobile Developer':    ['Kotlin','Swift','React Native','Firebase','CI/CD','AWS'],
                'DevOps Engineer':     ['Terraform','Ansible','GCP','Azure','Prometheus','Grafana'],
            }
            fallback_role = inferred_role if inferred_role in ROLE_GAP_MAP else 'Software Engineer'
            for sk in ROLE_GAP_MAP.get(fallback_role, []):
                if sk.lower() not in user_skills_lower:
                    gap_skills.append({'skill': sk, 'demand': 3, 'priority': 'high'})
            generic = ['Docker','AWS','Git','Agile','REST API','PostgreSQL','Redis','CI/CD']
            for sk in generic:
                if sk.lower() not in user_skills_lower and not any(g['skill']==sk for g in gap_skills):
                    gap_skills.append({'skill': sk, 'demand': 2, 'priority': 'medium'})
                if len(gap_skills) >= 10:
                    break
        # --- Fallback: static matches when scraping returns nothing ---
        if not top_jobs:
            top_jobs = [
                {
                    'title': 'Software Engineering Intern',
                    'company': 'TechNova Solutions',
                    'location': 'Remote',
                    'mode': 'Remote',
                    'snippet': 'We are looking for a Software Engineering Intern with experience in Python and JavaScript to build scalable microservices. Must be familiar with Agile methodologies.',
                    'url': '#',
                    'matchScore': 92,
                    'insight': {'tag': 'Startup', 'color': '#a855f7', 'desc': 'Startup culture - high ownership, fast-paced, direct impact'}
                },
                {
                    'title': 'Data Science Intern',
                    'company': 'Global Analytics',
                    'location': 'Bengaluru, India',
                    'mode': 'Hybrid',
                    'snippet': 'Join our analytics team! You will work with Pandas, Scikit-Learn, and SQL to extract insights from massive datasets.',
                    'url': '#',
                    'matchScore': 88,
                    'insight': {'tag': 'Tech Co.', 'color': '#6366f1', 'desc': 'Tech company - collaborative culture, learning opportunities'}
                },
                {
                    'title': 'Full Stack Developer Intern',
                    'company': 'Innovate Inc.',
                    'location': 'Pune, India',
                    'mode': 'On-site',
                    'snippet': 'Seeking a talented Full Stack Intern. You will work with React, Node.js, and MongoDB. Experience with cloud platforms (AWS) is a plus.',
                    'url': '#',
                    'matchScore': 85,
                    'insight': {'tag': 'Agile', 'color': '#00e4b8', 'desc': 'Agile shop - sprints, standups, iterative delivery'}
                },
                {
                    'title': 'Cloud DevOps Intern',
                    'company': 'CloudOps Systems',
                    'location': 'Remote',
                    'mode': 'Remote',
                    'snippet': 'Learn cloud infrastructure automation using Kubernetes, Docker, and CI/CD pipelines. Remote-first team.',
                    'url': '#',
                    'matchScore': 81,
                    'insight': {'tag': 'Remote-First', 'color': '#00f0ff', 'desc': 'Remote-first team - async culture, flexible hours'}
                }
            ]
            for job in top_jobs:
                if 'insight' not in job:
                    job['insight'] = _company_insight(job.get('snippet', ''), job.get('company', ''))

        return jsonify({
            "name": profile.get("fullName", ""),
            "email": profile.get("email", ""),
            "phone": profile.get("phone", ""),
            "location": profile.get("location", ""),
            "experience_years": profile.get("yearsExperience", ""),
            "linkedin": profile.get("linkedin", ""),
            "github": profile.get("github", ""),
            "portfolio": profile.get("portfolio", ""),
            "bio": profile.get("summary", ""),
            "college": profile.get("education", [])[0] if profile.get("education") else "",
            "skills": skills,
            "role": inferred_role,
            "stats": {
                "matchedInternships": matched_count,
                "interviewsCompleted": interviews_completed,
                "avgMatchScore": avg_score,
                "dayStreak": max(1, len(skills) // 4)
            },
            "topMatches": top_jobs,
            "skillGap": gap_skills,
            "marketTrends": market_trends,
            "session_id": (update_data.get('session_id') if 'update_data' in locals() else request.form.get("session_id")),
            "rawTextPreview": text[:500] + ("..." if len(text) > 500 else "")
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Backend crash during analysis: {str(e)}"}), 500


def _extract_company(url, snippet):
    """Try to extract company name from URL or snippet."""
    if 'linkedin.com' in url:
        return 'LinkedIn Listing'
    if 'indeed.com' in url:
        return 'Indeed Listing'
    if 'internshala.com' in url:
        return 'Internshala'
    if 'glassdoor.com' in url:
        return 'Glassdoor'
    # Try to pull from snippet
    for word in ['at ', 'with ', '@ ']:
        if word in snippet.lower():
            idx = snippet.lower().index(word) + len(word)
            return snippet[idx:idx+20].split()[0].strip(',.') if idx < len(snippet) else 'Company'
    return 'Tech Company'


def _guess_mode(snippet):
    s = snippet.lower()
    if 'remote' in s:
        return 'Remote'
    if 'hybrid' in s:
        return 'Hybrid'
    return 'On-site'


@app.route("/api/generate-cover-letter", methods=["POST"])
def generate_cover_letter():
    try:
        payload = request.get_json(silent=True) if request.is_json else {}
        if not isinstance(payload, dict):
            payload = {}

        job_title = (
            request.form.get("jobTitle", "").strip()
            or request.form.get("job_title", "").strip()
            or str(payload.get("jobTitle", "")).strip()
            or str(payload.get("job_title", "")).strip()
            or "Software Engineer"
        )
        company = (
            request.form.get("company", "").strip()
            or str(payload.get("company", "")).strip()
            or "the Company"
        )
        job_description = (
            request.form.get("jobDescription", "").strip()
            or request.form.get("job_description", "").strip()
            or str(payload.get("jobDescription", "")).strip()
            or str(payload.get("job_description", "")).strip()
            or ""
        )
        from datetime import date
        today = date.today().strftime("%B %d, %Y")

        stream = get_resume_stream_from_req(request)
        if not stream:
            return jsonify({"error": "No resume found."}), 400

        # If user uploads a resume in this request, persist it as their latest resume
        # so subsequent cover-letter generations can reuse it automatically.
        uploaded_resume = request.files.get("resume")
        if uploaded_resume and uploaded_resume.filename:
            user = _get_user_from_request()
            if user:
                try:
                    latest_session_id = str(uuid.uuid4())
                    latest_path = os.path.join(UPLOAD_FOLDER, f"{latest_session_id}.pdf")
                    uploaded_resume.seek(0)
                    uploaded_resume.save(latest_path)
                    db.update_user(user["_id"], session_id=latest_session_id)
                except Exception as ex:
                    print(f"[Cover Letter] Could not persist uploaded resume: {ex}")

        from mock_interview import extract_text_from_pdf
        try:
            resume_text = extract_text_from_pdf(stream)
        except Exception as e:
            return jsonify({"error": f"Failed to read resume: {str(e)}"}), 500

        if len(resume_text.strip()) < 30:
            return jsonify({"error": "Could not extract readable text from this file."}), 400

        # Build strict official-format prompt
        prompt_text = f"""You are a professional cover letter writer. Generate a complete, formal business cover letter.

Requirements:
- Use the exact date: {today}
- Target position: {job_title}
- Target company: {company}
- Use the candidate's details extracted from the resume below.

Resume Content:
---
{resume_text[:2500]}
---

Job Description (if provided):
---
{job_description[:1500] or "Not provided"}
---

The cover letter MUST follow this EXACT official format:

[Candidate Full Name]
[Candidate Email]
[Candidate Phone, if found in resume]
[City, if found in resume]
{today}

Hiring Manager
{company}
Subject: Application for the Position of {job_title}

Dear Hiring Manager,

[Opening paragraph: Express strong enthusiasm for the role and company. Mention where you learned about the position.]

[Body paragraph: Highlight 2-3 specific skills, experiences, or projects from the resume most relevant to {job_title}. Use concrete details.]

[Closing paragraph: Reiterate interest, mention availability for an interview, and thank them for their consideration.]

Sincerely,
[Candidate Full Name]

Return ONLY the formatted cover letter text, no extra explanations."""

        try:
            # Call configured LLM
            response = llm.invoke([HumanMessage(content=prompt_text)])
            cl_text = _normalize_llm_content(getattr(response, "content", ""))
            if not cl_text:
                raise ValueError("LLM returned empty cover letter content")
        except Exception as llm_error:
            print(f"[Cover Letter] LLM failed ({llm_error}), using fallback generator.")
            from mock_interview import extract_resume_profile
            profile = extract_resume_profile(resume_text)
            
            cand_name = profile.get("fullName") or "[Candidate Full Name]"
            cand_email = profile.get("email") or "[Candidate Email]"
            cand_phone = profile.get("phone") or "[Candidate Phone]"
            cand_location = profile.get("location") or "[Candidate Location]"
            skills = profile.get("skills", [])
            skills_str = ", ".join(skills[:3]) if skills else "relevant technical skills"
            
            cl_text = f"{cand_name}\n{cand_email}\n{cand_phone}\n{cand_location}\n{today}\n\nHiring Manager\n{company}\nSubject: Application for the Position of {job_title}\n\nDear Hiring Manager,\n\nI am writing to express my strong enthusiasm for the {job_title} position at {company}. With my background and strong foundation in {skills_str}, I am confident in my ability to contribute effectively to your team and make an immediate impact.\n\nThroughout my academic and professional journey, I have developed a skill set that aligns closely with the requirements for this role. I have consistently demonstrated a commitment to delivering high-quality work and quickly adapting to new technical challenges. My past projects have taught me the importance of problem-solving, efficient execution, and continuous learning.\n\nI am particularly drawn to {company} because of your innovative approach and industry leadership. I would welcome the opportunity to discuss how my background and skills will be beneficial to your organization. Thank you for considering my application. I look forward to the possibility of an interview.\n\nSincerely,\n{cand_name}"

        # Return plain text for editing (use cover_letter key as expected by frontend)
        return jsonify({"cover_letter": cl_text, "text": cl_text})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Failed to generate cover letter: {str(e)}"}), 500


@app.route("/api/download-cover-letter", methods=["POST"])
def download_cover_letter():
    """Generate PDF or DOCX from user-edited cover letter text."""
    try:
        data = request.json
        text = data.get("text", "").strip()
        fmt = data.get("format", "pdf").lower()

        if not text:
            return jsonify({"error": "No text provided"}), 400

        if fmt == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("helvetica", size=11)

            # Encode the whole text safely to latin-1, then write in one shot
            safe_text = text.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 6, safe_text)

            pdf_bytes = bytes(pdf.output())
            from flask import send_file
            return send_file(io.BytesIO(pdf_bytes), mimetype='application/pdf', as_attachment=True, download_name='Cover_Letter.pdf')

        elif fmt == "docx":
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = docx_lib.Document()

            # Set margins
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.2)
                section.right_margin = Inches(1.2)

            lines = text.split('\n')
            for i, line in enumerate(lines):
                if not line.strip():
                    doc.add_paragraph()
                elif i < 6:  # Header block — bold
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip())
                    run.bold = True
                    run.font.size = Pt(11)
                else:
                    p = doc.add_paragraph(line.strip())
                    p.runs[0].font.size = Pt(11) if p.runs else None

            docx_io = io.BytesIO()
            doc.save(docx_io)
            docx_io.seek(0)
            from flask import send_file
            return send_file(docx_io, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name='Cover_Letter.docx')

        else:
            return jsonify({"error": "Invalid format. Use 'pdf' or 'docx'"}), 400

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Download generation failed: {str(e)}"}), 500




@app.route("/contact")
def contact():
    return render_template("contact.html")

@app.route("/mock-interview/", defaults={"path": ""})
@app.route("/mock-interview/<path:path>")
def serve_mock_interview(path):
    return render_template("mock-interview.html")

@app.route("/documentation")
def documentation():
    return render_template("documentation.html")


@app.route("/privacy")
def privacy():
    return render_template("privacy.html")


@app.route("/scan", methods=["POST"])
def scan_jobs():
    """API endpoint for the 'Start AI Scan' button — returns real job results."""
    import random
    from tools import (scrape_linkedin, scrape_indeed,
                       scrape_freshershub, scrape_internshiphub, scrape_placementindia,
                       scrape_unstop, scrape_social_media)

    data = request.get_json() or {}
    query = data.get("query", "software intern India")
    if "india" not in query.lower():
        query = f"{query} India"

    all_cards = []
    seen_links = set()
    scrapers = [
        scrape_linkedin, scrape_indeed,
        scrape_freshershub, scrape_internshiphub, scrape_placementindia,
        scrape_unstop, scrape_social_media
    ]
    for func in scrapers:
        try:
            for card in func(query, days_ago=7):
                link = card.get("apply_link", "")
                if link and link not in seen_links:
                    seen_links.add(link)
                    all_cards.append(card)
        except Exception as e:
            print(f"[scan] scraper error: {e}")

    all_cards = filter_india_jobs(all_cards)

    # Sort by applicants ascending (lowest competition first)
    all_cards.sort(key=lambda x: x.get("applicants", 10000))
    for i, card in enumerate(all_cards):
        card["id"] = i + 1

    return jsonify({"jobs": all_cards})


def _normalize_llm_content(content):
    """Normalize provider-specific LLM content blocks into plain text."""
    if content is None:
        return ""
    if isinstance(content, str):
        return content.strip()
    if isinstance(content, dict):
        text = content.get("text") or content.get("content") or content.get("output_text") or ""
        return str(text).strip()
    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, str):
                text = item.strip()
                if text:
                    parts.append(text)
                continue
            if isinstance(item, dict):
                text = item.get("text") or item.get("content") or item.get("output_text") or ""
                if text:
                    parts.append(str(text).strip())
                continue
            text = getattr(item, "text", "") or getattr(item, "content", "")
            if text:
                parts.append(str(text).strip())
        return "\n".join(p for p in parts if p).strip()
    return str(content).strip()


@app.route("/chat", methods=["POST"])
def chat():
    """Handle chat messages from the user."""
    try:
        data = request.get_json(silent=True) or {}
        user_message = data.get("message", "").strip()
        session_id = data.get("session_id", "default")

        if not user_message:
            return jsonify({"error": "Empty message"}), 400

        # Get chat history
        get_chat_history(session_id)
        append_chat_messages(session_id, HumanMessage(content=user_message))
        history = get_chat_history_snapshot(session_id)

        # --- LIGHTNING-FAST TOOL INTERCEPTOR ---
        # Don't wait 15-30 seconds for the local CPU LLM to process if we already know they want jobs!
        search_keywords = ["intern", "job", "role", "developer", "engineer", "position", "work", "find", "show", "remote"]
        needs_search = any(k in user_message.lower() for k in search_keywords)
        
        ai_response = ""
        
        if needs_search:
            # ── SOURCE FILTER: detect if the user wants a specific platform only ──
            # Determine source filter
            msg_lower = user_message.lower()
            discovered_sources = []
            
            source_map = {
                "linkedin": "linkedin",
                "indeed": "indeed",
                "freshershub": "freshershub",
                "placementhub": "placementhub",
                "internshiphub": "internshiphub",
                "unstop": "unstop",
                "twitter": "social",
                "facebook": "social",
                "instagram": "social",
                "social": "social",
                "placementindia": "placementindia",
            }
            
            for kw, src in source_map.items():
                if kw in msg_lower:
                    discovered_sources.append(src)
            # Deduplicate (e.g. twitter + facebook both map to "social")
            discovered_sources = list(dict.fromkeys(discovered_sources))
                    
            sources_param = ",".join(discovered_sources) if discovered_sources else "linkedin,indeed,freshershub,placementhub,internshiphub,unstop,social"

            tool_output = search_internships.invoke({
                "query": user_message,
                "sources": sources_param
            })
            
            source_labels = discovered_sources if discovered_sources else ["All Sources"]
            source_title = "/".join(s.capitalize() for s in source_labels)
            
            if "```internship_cards" in tool_output:
                ai_response = f"Here are live jobs from **{source_title}**:\n\n" + tool_output
            else:
                ai_response = tool_output
        else:
            # Only run the LLM natively if it's conversational small talk
            llm_with_tools = llm.bind_tools(tools)
            ai_msg = llm_with_tools.invoke(history)
            
            if ai_msg.tool_calls:
                first_call = ai_msg.tool_calls[0] if ai_msg.tool_calls else {}
                if isinstance(first_call, dict):
                    call_args = first_call.get("args", {}) or {}
                else:
                    call_args = getattr(first_call, "args", {}) or {}
                query_arg = call_args.get("query", user_message) if isinstance(call_args, dict) else user_message
                tool_output = search_internships.invoke({"query": query_arg, "sources": "linkedin,indeed,freshershub,placementhub,internshiphub,unstop,social"})
                if "```internship_cards" in tool_output:
                    ai_response = "Here are the live job listings:\n\n" + tool_output
                else:
                    ai_response = tool_output
            else:
                ai_response = _normalize_llm_content(ai_msg.content)
                if ai_response.startswith("{") and ai_response.endswith("}"):
                    ai_response = "I couldn't quite understand that. Are you looking for a specific internship?"

        if not ai_response:
            ai_response = "I'm sorry, I couldn't process that. Could you try rephrasing? 😊"

        # Update history with AI's tool or text response
        if "```internship_cards" in ai_response:
            # Do NOT store 3,000 tokens of JSON in the LLM's chat history! It will freeze local models.
            append_chat_messages(session_id, AIMessage(content="I successfully found and displayed the matching internship cards to the user."))
        else:
            append_chat_messages(session_id, AIMessage(content=ai_response))

        return jsonify({
            "response": ai_response,
            "status": "success"
        })

    except Exception as e:
        print(f"Error in /chat: {e}")
        error_msg = str(e)
        if "connection error" in error_msg.lower() or "connectionrefusederror" in error_msg.lower():
            return jsonify({
                "response": "⚠️ **Ollama Error**: Could not connect to local Ollama instance. Please ensure Ollama is running and you have run `ollama run mistral`.",
                "status": "success"
            })
        return jsonify({
            "error": f"Something went wrong: {str(e)}",
            "status": "error"
        }), 500


@app.route("/upload-resume", methods=["POST"])
def upload_resume():
    """
    Handle PDF resume uploads:
    1. Extract text via PyPDF2
    2. Parse key skills with a fast local regex extractor (NO LLM = instant)
    3. Search all platforms for jobs posted in the LAST 7 DAYS
    4. Return job cards + a skill summary
    """
    import re

    try:
        stream = get_resume_stream_from_req(request)
        if not stream:
            return jsonify({"error": "No resume found"}), 400

        session_id = request.form.get("session_id", "default")

        # ── Step 1: Extract text from PDF ─────────────────────────────────
        reader = PdfReader(stream)
        resume_text = ""
        for page in reader.pages:
            text = page.extract_text()
            if text:
                resume_text += text + "\n"

        if not resume_text.strip():
            return jsonify({"error": "Could not extract text from the PDF. Please try a text-based PDF."}), 400

        # ── Step 2: Fast keyword extractor (NO LLM — instant) ─────────────
        text_lower = resume_text.lower()

        TECH_SKILLS = [
            "python", "javascript", "typescript", "java", "c++", "c#", "go", "rust",
            "react", "angular", "vue", "node.js", "nodejs", "django", "flask", "fastapi",
            "spring", "express", "next.js", "nextjs", "sql", "mysql", "postgresql",
            "mongodb", "redis", "machine learning", "deep learning", "nlp",
            "computer vision", "data science", "tensorflow", "pytorch", "keras",
            "scikit-learn", "pandas", "numpy", "aws", "azure", "gcp", "docker",
            "kubernetes", "git", "html", "css", "rest api", "graphql",
            "android", "ios", "flutter", "react native", "swift", "kotlin",
            "blockchain", "solidity", "web3", "devops"
        ]

        ROLE_KEYWORDS = [
            "software engineer", "software developer", "backend developer",
            "frontend developer", "full stack developer", "data scientist",
            "data analyst", "ml engineer", "machine learning engineer",
            "ai engineer", "devops engineer", "android developer", "ios developer",
            "mobile developer", "web developer", "intern", "fresher"
        ]

        found_skills = [s for s in TECH_SKILLS if s in text_lower]
        found_roles = [r for r in ROLE_KEYWORDS if r in text_lower]

        # ── Step 3: Build ROTATING query variations ────────────────────────
        # Shuffle skills so every upload picks different combos as primary keys
        import random
        shuffled_skills = found_skills[:] if found_skills else ["software"]
        random.shuffle(shuffled_skills)

        primary_role = found_roles[0] if found_roles else "developer"

        # Build up to 3 unique queries from different skill combos
        queries = []
        if len(shuffled_skills) >= 2:
            queries.append(f"{shuffled_skills[0]} {primary_role} intern india")
            queries.append(f"{shuffled_skills[1]} developer intern india")
        if len(shuffled_skills) >= 3:
            queries.append(f"{shuffled_skills[2]} engineer internship india")
        if not queries:
            queries = [f"{shuffled_skills[0]} {primary_role} intern india"]

        search_query = queries[0]  # primary query for display

        # ── Step 4: Multi-query scraping with STRICT 7-day filter ──────────
        from tools import scrape_linkedin, scrape_indeed, \
                          scrape_freshershub, scrape_internshiphub, scrape_placementindia, \
                          scrape_unstop, scrape_social_media

        seen_links = set()
        all_cards = []

        # Run all generated queries to get broad variety
        scrapers = [
            scrape_linkedin, scrape_indeed, 
            scrape_freshershub, scrape_internshiphub, scrape_placementindia,
            scrape_unstop, scrape_social_media
        ]

        for q in queries:
            for scraper_func in scrapers:
                try:
                    # Fetch fresh jobs (7 days) for this query variation
                    results = scraper_func(q, days_ago=7)
                    for card in results:
                        link = card.get("apply_link", "")
                        if link and link not in seen_links:
                            seen_links.add(link)
                            all_cards.append(card)
                except Exception as e:
                    print(f"Scraper error for query '{q}': {e}")

        all_cards = filter_india_jobs(all_cards)

        # Shuffle the final pool to ensure a different visual order every time
        random.shuffle(all_cards)

        # Re-number and badge (after shuffle)
        for i, card in enumerate(all_cards):
            card["id"] = i + 1
            source = card.pop("source", "Job Board")
            card["title"] = f"[{source}] {card['title']}"

        # ── Step 4: Build the response ─────────────────────────────────────
        skills_display = ", ".join(f"**{s.title()}**" for s in found_skills[:8]) or "General Software Skills"
        roles_display = ", ".join(r.title() for r in found_roles[:3]) or "Software Developer"

        if all_cards:
            cards_block = "```internship_cards\n" + json.dumps(all_cards, indent=2) + "\n```"
            intro = (
                f"📄 **Resume Analyzed!**\n\n"
                f"🔑 **Skills Detected:** {skills_display}\n"
                f"💼 **Target Roles:** {roles_display}\n"
                f"🔍 **Searched for:** `{search_query}`\n"
                f"⏰ **Filter:** Jobs posted in the **last 7 days** only\n\n"
                f"Found **{len(all_cards)} fresh, open opportunities** matching your profile:\n\n"
            )
            ai_response = intro + cards_block
        else:
            ai_response = (
                f"📄 **Resume Analyzed!**\n\n"
                f"🔑 **Skills Detected:** {skills_display}\n"
                f"💼 **Target Roles:** {roles_display}\n\n"
                f"⚠️ No jobs found in the **last 7 days** for `{search_query}`. "
                f"Ask me to search with a wider range, e.g. *'find python developer jobs from last month'*."
            )

        append_chat_messages(session_id, [
            HumanMessage(content=f"[Resume analyzed — skills: {', '.join(found_skills[:5])}]"),
            AIMessage(content=ai_response)
        ])

        return jsonify({
            "response": ai_response,
            "status": "success",
            "skills_found": found_skills,
            "roles_found": found_roles,
            "search_query": search_query,
            "jobs_count": len(all_cards)
        })

    except Exception as e:
        print(f"Error in /upload-resume: {e}")
        return jsonify({"error": f"Failed to process resume: {str(e)}", "status": "error"}), 500



@app.route("/api/avatar/session", methods=["POST"])
def create_avatar_session():
    """Create a new Akool Live Avatar streaming session."""
    akool_api_key = os.getenv("AKOOL_API_KEY")
    if not akool_api_key or akool_api_key == "your_akool_api_key_here":
        return jsonify({"error": "AKOOL_API_KEY is not set in the .env file"}), 500

    url = "https://openapi.akool.com/api/v3/streamingAvatar/session/create"
    headers = {
        "Authorization": f"Bearer {akool_api_key}",
        "Content-Type": "application/json"
    }
    
    # User should set AKOOL_AVATAR_ID in .env, fallback to a commonly used default format or name
    avatar_id = os.getenv("AKOOL_AVATAR_ID", "default_avatar") 
    
    payload = {
        "avatar_id": avatar_id
    }

    try:
        response = requests.post(url, headers=headers, json=payload, timeout=EXTERNAL_HTTP_TIMEOUT)
        response.raise_for_status()
        data = response.json()
        return jsonify(data)
    except requests.exceptions.HTTPError as e:
        print(f"HTTP Error creating Akool session: {e.response.text}")
        return jsonify({"error": "Failed to create avatar session", "details": e.response.text}), e.response.status_code
    except Exception as e:
        print(f"Error creating Akool session: {e}")
        return jsonify({"error": str(e)}), 500



# ──────────────────────────────────────────────
# Mock Interview Question Generator
# ──────────────────────────────────────────────
@app.route('/api/interview/generate-questions', methods=['POST'])
def generate_interview_questions():
    try:
        data = request.get_json() or {}
        difficulty = data.get('difficulty', 'medium')
        try:
            count = int(data.get('count', 20))
        except (TypeError, ValueError):
            count = 20
        from mock_interview import generate_real_life_questions, infer_role

        incoming_profile = data.get("profile")
        if isinstance(incoming_profile, dict) and incoming_profile:
            profile = dict(incoming_profile)
        else:
            skills = data.get("skills", []) or []
            role = data.get("role") or (infer_role(skills) if skills else "Software Engineer")
            profile = {
                "role": role,
                "skills": skills,
                "projects": [],
                "experienceHighlights": [],
                "achievements": [],
                "yearsExperience": 0
            }

        questions = generate_real_life_questions(profile, difficulty=difficulty, count=count)
        return jsonify({
            'questions': questions,
            'difficulty': difficulty,
            'count': len(questions),
            'role': profile.get("role", "Software Engineer")
        })
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/fast-extract-skills', methods=['POST'])
def fast_extract_skills():
    if 'resume' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    try:
        file = request.files['resume']
        from mock_interview import extract_text_from_pdf, extract_resume_profile
        text = extract_text_from_pdf(file.stream)
        if not text or len(text) < 30:
            return jsonify({"error": "Could not extract readable text from the resume PDF"}), 400
        profile = extract_resume_profile(text)
        return jsonify({
            "extractedSkills": profile.get("skills", []),
            "role": profile.get("role", "Software Engineer"),
            "profile": profile
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ──────────────────────────────────────────────
# Run the server
# ──────────────────────────────────────────────
if __name__ == "__main__":
    setup_llm()
    setup_agent()
    host = os.getenv("HOST", "0.0.0.0")
    port = int(os.getenv("PORT", "5000"))
    debug_flag = os.getenv("DEBUG", "").strip().lower() in {"1", "true", "yes"}
    env_name = os.getenv("FLASK_ENV", "").strip().lower()
    is_dev = debug_flag or env_name == "development"

    print("[START] AI Internship Finder Agent is running!")
    print(f"[INFO] Listening on http://{host}:{port}")
    print("")

    if not is_dev:
        try:
            from waitress import serve
            threads = max(4, int(os.getenv("WAITRESS_THREADS", "16")))
            connection_limit = max(threads * 4, int(os.getenv("WAITRESS_CONNECTION_LIMIT", str(threads * 8))))
            print(f"[PROD] Waitress enabled with threads={threads}, connection_limit={connection_limit}")
            serve(app, host=host, port=port, threads=threads, connection_limit=connection_limit)
        except Exception as ex:
            print(f"[WARN] Waitress unavailable ({ex}). Falling back to Flask threaded server.")
            app.run(debug=False, host=host, port=port, threaded=True, use_reloader=False)
    else:
        print("[DEV] Running Flask debug server")
        app.run(debug=True, host=host, port=port, threaded=True, use_reloader=True)
print("[FINISH] app.py loaded")
